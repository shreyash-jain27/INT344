# ============================================
# app.py - Simple Web Interface for Day 2
# ============================================

import sys
from pathlib import Path
import streamlit as st
import time
import io

# Add src to path
sys.path.append(str(Path(__file__).parent / "src"))

# --- Optimized Imports ---
# Heavy AI models are now loaded lazily inside functions



# Page configuration
st.set_page_config(
    page_title="Annual Report Analyzer",
    page_icon="📊",
    layout="wide"
)

# Standardize Plotly Theme
try:
    import plotly.io as pio
    pio.templates.default = "plotly_white"
except:
    pass

# Custom CSS
st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');
    
    html, body, [class*="css"] {
        font-family: 'Inter', sans-serif;
    }
    
    .big-font {
        font-size: 28px !important;
        font-weight: 700;
        color: #1E293B;
    }
    
    .section-box {
        background-color: #ffffff;
        padding: 24px;
        border-radius: 12px;
        border: 1px solid #E2E8F0;
        box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
        margin: 16px 0;
    }
    
    .stat-box {
        background: linear-gradient(135deg, #6366F1 0%, #4F46E5 100%);
        color: white;
        padding: 20px;
        border-radius: 12px;
        text-align: center;
        box-shadow: 0 10px 15px -3px rgba(79, 70, 229, 0.2);
    }
    
    .summary-card {
        background-color: #f8fafc;
        border-left: 5px solid #6366F1;
        padding: 20px;
        border-radius: 0 12px 12px 0;
        margin: 10px 0;
        font-size: 1.1rem;
        line-height: 1.8;
        color: #334155;
    }
    
    .metadata-tag {
        display: inline-block;
        padding: 4px 12px;
        background-color: #E2E8F0;
        border-radius: 20px;
        font-size: 0.85rem;
        font-weight: 600;
        color: #475569;
        margin-right: 8px;
    }

    .premium-header {
        background: linear-gradient(90deg, #4F46E5 0%, #7C3AED 100%);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        font-weight: 800;
    }
    </style>
""", unsafe_allow_html=True)

# Title and description
st.title("📊 Annual Report Analyzer (Advanced)")
st.markdown("**An intelligent system to extract, analyze, and summarize Indian annual reports with multilingual support.**")
st.markdown("---")

# Sidebar
with st.sidebar:
    st.header("📖 About")
    st.info("""
    This tool extracts and analyzes annual reports:
    
    ✓ Text extraction from PDFs
    ✓ Automatic section identification
    ✓ Table extraction
    ✓ Statistical analysis
    """)

    st.header("🎯 Features")
    st.markdown("""
    - **Chairman's Message**
    - **Financial Performance**
    - **Business Overview**
    - **Future Outlook**
    - **And more...**
    """)

    st.header("📝 Instructions")
    st.markdown("""
    1. Upload a PDF file
    2. Wait for processing
    3. Explore extracted sections
    4. Download results
    """)

# Main content
col1, col2 = st.columns([2, 1])

with col1:
    st.subheader("📁 Upload PDF")
    uploaded_file = st.file_uploader(
        "Choose an annual report PDF",
        type="pdf",
        help="Upload a PDF file of an annual report"
    )

with col2:
    st.subheader("📊 Quick Stats")
    if uploaded_file:
        file_size = len(uploaded_file.getvalue()) / (1024 * 1024)  # MB
        st.metric("File Size", f"{file_size:.2f} MB")
    else:
        st.info("Upload a file to see stats")

# Process the uploaded file
if uploaded_file:
    # Save temporary file
    temp_path = Path("temp_upload.pdf")
    with open(temp_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    # --- OPTIMIZATION START ---
    
    @st.cache_resource
    def get_pdf_processor(file_path):
        from src.pdf_processor import PDFProcessor
        return PDFProcessor(file_path)

    @st.cache_data
    def process_document(_processor):
        try:
            with st.status("🔍 Document Analysis Engine") as status:
                status.write("📄 Extracting text and sections...")
                text = _processor.extract_text()
                sections = _processor.identify_sections()
                section_stats = _processor.get_section_stats()
                
                status.write("📊 Extracting tabular data...")
                table_results = _processor.extract_tables_detailed()
                
                status.write("🎨 Analyzing visual layout...")
                from src.layout_analyzer import analyze_pdf_layout
                layout_results = analyze_pdf_layout(_processor.pdf_path)
                
                status.update(label="✅ Analysis Complete!", state="complete")
                
            return {
                'text': text,
                'sections': sections,
                'section_stats': section_stats,
                'table_results': table_results,
                'layout_results': layout_results,
                'metadata': _processor.metadata
            }
        except Exception as e:
            st.error(f"🚨 Critical error during processing: {str(e)}")
            st.stop()

    # Clear session state if a new file is uploaded
    if "current_file" not in st.session_state or st.session_state.current_file != uploaded_file.name:
        import gc
        # Clear large state objects
        st.session_state.clear()
        st.session_state.current_file = uploaded_file.name
        # Manually trigger garbage collection
        gc.collect()

    # Run processing once
    if 'doc_data' not in st.session_state:
        processor = get_pdf_processor(str(temp_path))
        st.session_state['doc_data'] = process_document(processor)
        st.success("✅ Document analyzed and ready!")

    # Get data from session state
    doc_data = st.session_state['doc_data']
    text = doc_data['text']
    sections = doc_data['sections']
    section_stats = doc_data['section_stats']
    table_results = doc_data['table_results']
    layout_results = doc_data['layout_results']
    metadata = doc_data['metadata']
    processor = get_pdf_processor(str(temp_path)) # Cached
    # --- OPTIMIZATION END ---

    # Create tabs for different views
    tabs = st.tabs(["📄 Overview", "📑 Sections", "📊 Tables", "📈 Visuals", "🎨 Layout", "📝 Summaries", "🌍 Multilingual", "💾 Download"])

    with tabs[0]:
        st.header("Document Overview")

        # Display metadata
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown('<div class="stat-box">', unsafe_allow_html=True)
            st.metric("Total Pages", metadata.get('total_pages', 0))
            st.markdown('</div>', unsafe_allow_html=True)

        with col2:
            st.markdown('<div class="stat-box">', unsafe_allow_html=True)
            st.metric("Total Characters", f"{len(text):,}")
            st.markdown('</div>', unsafe_allow_html=True)

        with col3:
            st.markdown('<div class="stat-box">', unsafe_allow_html=True)
            st.metric("Total Words", f"{len(text.split()):,}")
            st.markdown('</div>', unsafe_allow_html=True)

        # Text preview
        with st.expander("📖 View Full Text (First 2000 characters)"):
            st.text_area("Extracted Text", text[:2000] + "...", height=300)

    with tabs[1]:
        st.header("Identified Sections")

        if sections:
            st.success(f"✅ Found {len(sections)} sections")

            # Display each section
            for i, (section_name, content) in enumerate(sections.items(), 1):
                with st.expander(f"**{i}. {section_name}**", expanded=(i == 1)):
                    # Section stats from doc_data
                    stats = section_stats[section_name]
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric(
                            "Words", f"{stats['word_count']:,}")
                    with col2:
                        st.metric("Characters",
                                  f"{stats['char_count']:,}")
                    with col3:
                        st.metric(
                            "Lines", stats['line_count'])

                    st.markdown("---")

                    # Display content
                    st.markdown("**Content:**")
                    st.text_area(
                        f"content_{i}",
                        content,
                        height=200,
                        label_visibility="collapsed"
                    )
            else:
                st.warning(
                    "⚠️ No sections identified. The PDF might have an unusual structure.")

    with tabs[3]:  # Visuals tab
        st.header("📈 Document Visualizations")
        
        if sections:
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("📊 Section Length Distribution")
                try:
                    import pandas as pd
                    import plotly.express as px
                    
                    sec_data = [{'Section': n, 'Words': len(c.split())} for n, c in sections.items()]
                    sec_df = pd.DataFrame(sec_data)
                    
                    fig = px.bar(
                        sec_df, x='Section', y='Words', color='Words',
                        title='Word Count by Section', color_continuous_scale='Blues'
                    )
                    st.plotly_chart(fig, use_container_width=True)
                except Exception as e:
                    st.error(f"Error generating chart: {e}")
                
            with col2:
                st.subheader("☁️ Section Word Cloud")
                selected_sec = st.selectbox("Select section:", list(sections.keys()))
                
                if st.button("Generate Word Cloud"):
                    try:
                        from wordcloud import WordCloud
                        import matplotlib.pyplot as plt
                        
                        with st.spinner("Generating..."):
                            wc = WordCloud(width=800, height=400, background_color='white', max_words=100).generate(sections[selected_sec])
                            fig_wc, ax = plt.subplots(figsize=(10, 5))
                            ax.imshow(wc, interpolation='bilinear')
                            ax.axis('off')
                            st.pyplot(fig_wc)
                    except Exception as e:
                        st.error(f"Error generating word cloud: {e}")
            
            st.markdown("---")
            st.subheader("🔝 Top Keywords (Overall)")
            # Simple keyword extraction from all text
            all_text = " ".join(sections.values())
            words = all_text.lower().split()
            # Filter short words
            words = [w for w in words if len(w) > 3]
            from collections import Counter
            common_words = Counter(words).most_common(15)
            
            keyword_df = pd.DataFrame(common_words, columns=['Keyword', 'Frequency'])
            fig_keywords = px.bar(
                keyword_df, 
                x='Frequency', 
                y='Keyword',
                orientation='h',
                title='Top 15 Most Frequent Keywords',
                color='Frequency',
                color_continuous_scale='Viridis'
            )
            st.plotly_chart(fig_keywords, use_container_width=True)
            
        else:
            st.info("Upload and process a document to see visualizations.")

    with tabs[4]:  # Layout Analysis tab (was 3)
        st.header("Layout Analysis")
        
        # Get results from session state
        blocks = layout_results['blocks']
        structure = layout_results['structure']
        stats = layout_results['stats']
        analyzer = layout_results['analyzer']
            
        # Statistics
        st.success(f"✅ Analyzed {stats['total_blocks']} layout blocks")
        
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            st.metric("Total Blocks", stats['total_blocks'])
        with col2:
            st.metric("Headings", len(structure['headings']))
        with col3:
            st.metric("Paragraphs", len(structure['paragraphs']))
        with col4:
            st.metric("Pages", metadata.get('total_pages', 0))
        
        st.markdown("---")
        
        # Block type distribution
        st.subheader("📊 Block Type Distribution")
        
        # Create pie chart
        type_data = pd.DataFrame([
            {'Type': k, 'Count': v}
            for k, v in stats['by_type'].items()
        ])
        
        fig = px.pie(
            type_data,
            values='Count',
            names='Type',
            title='Distribution of Block Types'
        )
        st.plotly_chart(fig, use_container_width=True)
        
        st.markdown("---")
        
        # Document structure
        st.subheader("📋 Document Structure")
        
        # Headings
        if structure['headings']:
            with st.expander(f"**Headings** ({len(structure['headings'])} found)", expanded=True):
                for i, heading in enumerate(structure['headings'], 1):
                    st.markdown(f"**{i}.** Page {heading['page']}: {heading['text']}")
        
        # Paragraphs summary
        if structure['paragraphs']:
            with st.expander(f"**Paragraphs** ({len(structure['paragraphs'])} found)"):
                para_df = pd.DataFrame([
                    {
                        'Paragraph': i,
                        'Page': p['page'],
                        'Words': p['word_count'],
                        'Preview': p['text']
                    }
                    for i, p in enumerate(structure['paragraphs'][:10], 1)
                ])
                st.dataframe(para_df, use_container_width=True, hide_index=True)
                
                if len(structure['paragraphs']) > 10:
                    st.info(f"Showing first 10 of {len(structure['paragraphs'])} paragraphs")
        
        st.markdown("---")
        
        # Layout visualization
        st.subheader("🎨 Layout Visualization")
        
        page_to_viz = st.selectbox(
            "Select page to visualize:",
            range(1, metadata.get('total_pages', 0) + 1),
            key="layout_viz_page"
        )
        
        if st.button("Generate Visualization"):
            with st.spinner("Creating visualization..."):
                viz_path = analyzer.visualize_layout(
                    page_num=page_to_viz,
                    output_path=f"layout_page_{page_to_viz}.png"
                )
                
                if viz_path and Path(viz_path).exists():
                    st.image(viz_path, caption=f"Layout Analysis - Page {page_to_viz}")
                else:
                    st.warning("Could not generate visualization")
        
        st.markdown("---")
            
        # Reading order
        st.subheader("📖 Reading Order")
        
        with st.expander("View blocks in reading order"):
            page_filter = st.selectbox(
                "Filter by page:",
                ["All"] + list(range(1, metadata.get('total_pages', 0) + 1)),
                key="layout_reading_filter"
            )
            
            if page_filter == "All":
                display_blocks = blocks[:20]  # First 20 blocks
            else:
                display_blocks = [b for b in blocks if b.page == page_filter][:20]
            
            for i, block in enumerate(display_blocks, 1):
                st.text_area(
                    f"{i}. {block.block_type.upper()} (Page {block.page})",
                    block.text,
                    height=100,
                    key=f"block_{i}"
                )
            
            if len(display_blocks) < len(blocks):
                st.info(f"Showing {len(display_blocks)} of {len(blocks)} total blocks")

    with tabs[5]:  # Summaries tab (was 4)
        st.header("AI-Powered Summaries")
        
        # Summarization controls
        col1, col2, col3 = st.columns(3)
        
        with col1:
            summ_method = st.selectbox(
                "Summarization Method",
                ["Hybrid (Best)", "Extractive (Fast)", "Abstractive (Quality)"],
                help="Hybrid combines both methods for best results"
            )
        
        with col2:
            compression = st.slider(
                "Compression Level",
                min_value=0.1,
                max_value=0.8,
                value=0.3,
                step=0.1,
                help="How much to compress (0.3 = 70% reduction)"
            )
        
        with col3:
            use_gpu = st.checkbox(
                "Use GPU Acceleration",
                value=True,
                help="Use your RTX 3060 for faster processing"
            )
        
        # Generate summaries button
        if st.button("🚀 Generate Summaries", type="primary"):
            
            with st.spinner("🤖 AI is analyzing and summarizing..."):
                from summarizer import UnifiedSummarizer
                
                # Map UI choice to method
                method_map = {
                    "Hybrid (Best)": "hybrid",
                    "Extractive (Fast)": "extractive",
                    "Abstractive (Quality)": "abstractive"
                }
                method = method_map[summ_method]
                
                # Initialize summarizer
                summarizer = UnifiedSummarizer(use_gpu=use_gpu)
                
                # Generate summaries
                start_time = time.time()
                results = summarizer.summarize_sections(
                    sections,
                    method=method
                )
                total_time = time.time() - start_time
                
                # Store in session state
                st.session_state['summaries'] = results
                st.session_state['summ_method'] = method
                st.session_state['summ_time'] = total_time
            
            st.success(f"✅ Summaries generated in {total_time:.1f}s!")
        
        # Display summaries if generated
        if 'summaries' in st.session_state:
            results = st.session_state['summaries']
            
            st.markdown("---")
            
            # Statistics
            st.subheader("📊 Summary Statistics")
            
            total_orig = sum(r.original_length for r in results.values())
            total_summ = sum(r.summary_length for r in results.values())
            avg_compression = (1 - total_summ/total_orig) * 100 if total_orig > 0 else 0
            
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Original Words", f"{total_orig:,}")
            with col2:
                st.metric("Summary Words", f"{total_summ:,}")
            with col3:
                st.metric("Compression", f"{avg_compression:.1f}%")
            with col4:
                st.metric("Processing Time", f"{st.session_state['summ_time']:.1f}s")
            
            st.markdown("---")
            
            # Display each summary
            st.subheader("📝 Section Summaries")
            
            for i, (section_name, result) in enumerate(results.items(), 1):
                with st.expander(
                    f"**{i}. {section_name}** "
                    f"({result.original_length} → {result.summary_length} words)",
                    expanded=(i <= 2)
                ):
                    # Summary stats
                    col1, col2, col3 = st.columns(3)
                    with col1:
                        st.metric("Original", f"{result.original_length} words")
                    with col2:
                        st.metric("Summary", f"{result.summary_length} words")
                    with col3:
                        st.metric("Compression", f"{result.compression_ratio*100:.1f}%")
                    
                    st.markdown("---")
                    
                    # Tabs for original vs summary
                    tab1, tab2 = st.tabs(["📝 Summary", "📄 Original"])
                    
                    with tab1:
                        st.markdown(f"""
                        <div class="summary-card">
                            {result.summary}
                        </div>
                        """, unsafe_allow_html=True)
                        
                        # Copy button in expander to keep UI clean
                        with st.expander("📋 Copy Summary"):
                            st.code(result.summary, language=None)
                    
                    with tab2:
                        st.markdown("**Original Text:**")
                        original_text = sections[section_name]
                        st.text_area(
                            "Original",
                            original_text,
                            height=300,
                            label_visibility="collapsed",
                            key=f"orig_{i}"
                        )
                    
                    # Quality metrics
                    with st.expander("📊 Quality Metrics"):
                        metrics = summarizer.evaluate_summary(
                            sections[section_name],
                            result.summary
                        )
                        
                        col1, col2 = st.columns(2)
                        with col1:
                            st.metric(
                                "Readability Score",
                                f"{metrics['flesch_reading_ease']:.1f}",
                                help="Higher is easier to read (60-70 is good)"
                            )
                        with col2:
                            st.metric(
                                "Grade Level",
                                f"{metrics['flesch_kincaid_grade']:.1f}",
                                help="U.S. school grade level"
                            )
                        
                        col3, col4 = st.columns(2)
                        with col3:
                            st.metric(
                                "Number Retention",
                                f"{metrics['number_retention']*100:.0f}%",
                                help="Percentage of numbers preserved"
                            )
                        with col4:
                            st.metric(
                                "Entity Retention",
                                f"{metrics['entity_retention']*100:.0f}%",
                                help="Percentage of names/entities preserved"
                            )
            
            # Download all summaries
            st.markdown("---")
            st.subheader("📥 Download Summaries")
            
            # Create downloadable text
            download_text = f"""
ANNUAL REPORT SUMMARY
{'=' * 80}

File: {uploaded_file.name}
Method: {st.session_state['summ_method']}
Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}

Total Compression: {avg_compression:.1f}%
Processing Time: {st.session_state['summ_time']:.1f}s

{'=' * 80}

"""
            
            for section_name, result in results.items():
                download_text += f"\n{'=' * 80}\n"
                download_text += f"{section_name.upper()}\n"
                download_text += f"{'=' * 80}\n\n"
                download_text += f"Original: {result.original_length} words | "
                download_text += f"Summary: {result.summary_length} words | "
                download_text += f"Compression: {result.compression_ratio*100:.1f}%\n\n"
                download_text += result.summary + "\n"
            
            st.download_button(
                label="📥 Download All Summaries (TXT)",
                data=download_text,
                file_name=f"{Path(uploaded_file.name).stem}_summaries.txt",
                mime="text/plain"
            )
            
            # Individual downloads
            cols = st.columns(3)
            for i, (section_name, result) in enumerate(results.items()):
                with cols[i % 3]:
                    st.download_button(
                        label=f"📄 {section_name[:20]}...",
                        data=result.summary,
                        file_name=f"{section_name.replace(' ', '_')}_summary.txt",
                        mime="text/plain",
                        key=f"dl_summ_{i}"
                    )
        
        else:
            st.info("👆 Click 'Generate Summaries' to start AI summarization")
            
            # Show what will be summarized
            if sections:
                st.markdown("### 📋 Sections to Summarize:")
                for i, (name, content) in enumerate(sections.items(), 1):
                    words = len(content.split())
                    st.markdown(f"{i}. **{name}** - {words:,} words")

    with tabs[2]:
        st.header("Extracted Tables")

        # Get results from session state
        tables = table_results['tables']
        stats = table_results['stats']

        if tables:
            # Display statistics
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Total Tables", stats['total_tables'])
            with col2:
                st.metric("Total Rows", stats['total_rows'])
            with col3:
                st.metric("With Headers", stats['tables_with_headers'])
            with col4:
                if stats['largest_table']:
                    st.metric("Largest Table",
                              f"{stats['largest_table']['rows']}×{stats['largest_table']['columns']}")

            st.success(f"✅ Found {len(tables)} tables")

            # Display each table
            for i, table in enumerate(tables, 1):
                with st.expander(
                    f"**Table {table['table_number']}** - Page {table['page']} "
                    f"({table['rows']}×{table['columns']})",
                    expanded=(i <= 3)  # Expand first 3 tables
                ):
                    # Display as DataFrame if available
                    if table.get('dataframe') is not None:
                        st.dataframe(
                            table['dataframe'],
                            use_container_width=True,
                            height=min(
                                400, len(table['dataframe']) * 35 + 50)
                        )

                        # Download option for this table
                        csv = table['dataframe'].to_csv(index=False)
                        st.download_button(
                            label=f"📥 Download Table {table['table_number']} as CSV",
                            data=csv,
                            file_name=f"table_{table['table_number']}_page_{table['page']}.csv",
                            mime="text/csv",
                            key=f"table_download_{i}"
                        )
                    else:
                        # Display as raw data
                        st.warning(
                            "⚠️ Could not parse as structured table. Showing raw data:")
                        # Show first 10 rows
                        for row_idx, row in enumerate(table['data'][:10]):
                            st.text(" | ".join(str(cell) for cell in row))

                        if len(table['data']) > 10:
                            st.info(
                                f"... and {len(table['data']) - 10} more rows")

            # Download all tables
            st.markdown("---")
            st.subheader("📥 Download All Tables")

            # Combine all tables into one Excel file
            try:
                from io import BytesIO
                import pandas as pd

                output = BytesIO()
                with pd.ExcelWriter(output, engine='openpyxl') as writer:
                    for table in tables:
                        if table.get('dataframe') is not None:
                            sheet_name = f"Table_{table['table_number']}_P{table['page']}"
                            table['dataframe'].to_excel(
                                writer,
                                # Excel sheet name limit
                                sheet_name=sheet_name[:31],
                                index=False
                            )

                st.download_button(
                    label="📥 Download All Tables (Excel)",
                    data=output.getvalue(),
                    file_name=f"{Path(uploaded_file.name).stem}_all_tables.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                )
            except ImportError:
                st.info(
                    "💡 Install openpyxl to enable Excel download: pip install openpyxl")

        else:
            st.info("ℹ️ No tables found in this document.")
            st.markdown("""
            **Possible reasons:**
            - Document doesn't contain tables
            - Tables are embedded as images
            - Table structure is too complex
            
            **Solutions:**
            - Try adjusting table extraction settings
            - Use OCR for image-based tables (Day 3)
            """)

    with tabs[6]:  # Multilingual tab (was 5)
        st.header("🌍 Multilingual Analysis")
        st.markdown("**Process and summarize in multiple Indian languages**")
        
        # Language detection results
        st.subheader("📊 Language Detection")
        
        # Cache language detection to prevent constant re-processing
        @st.cache_data
        def cached_detect_languages(_sections):
            from src.language_detector import LanguageDetector
            _detector = LanguageDetector()
            _lang_results = _detector.detect_document_language(_sections)
            _doc_lang, _doc_conf = _detector.get_document_primary_language(_sections)
            return _lang_results, _doc_lang, _doc_conf

        with st.spinner("Detecting languages..."):
            from src.language_detector import LanguageDetector
            detector = LanguageDetector()
            lang_results, doc_lang, doc_conf = cached_detect_languages(sections)
            
            # Display results
            if lang_results:
                lang_info = detector.get_language_info(doc_lang)
                
                col1, col2, col3 = st.columns(3)
                with col1:
                    st.metric(
                        "Primary Language",
                        lang_info['name'],
                        help="Main language of the document"
                    )
                with col2:
                    st.metric(
                        "Script",
                        lang_info['script'],
                        help="Writing system used"
                    )
                with col3:
                    st.metric(
                        "Confidence",
                        f"{doc_conf*100:.1f}%",
                        help="Detection confidence"
                    )
                
                st.markdown("---")
                
                # Language breakdown by section
                st.subheader("📑 Language by Section")
                
                lang_data = []
                for section_name, result in lang_results.items():
                    lang_data.append({
                        'Section': section_name,
                        'Language': result['language_name'],
                        'Script': result['script'],
                        'Confidence': f"{result['confidence']*100:.1f}%",
                        'Code-Mixed': '⚠️' if result['is_code_mixed'] else '✓'
                    })
                
                import pandas as pd
                df = pd.DataFrame(lang_data)
                st.dataframe(df, use_container_width=True, hide_index=True)
                
                # Code-mixed sections warning
                code_mixed_sections = [
                    name for name, res in lang_results.items()
                    if res.get('is_code_mixed', False)
                ]
                
                if code_mixed_sections:
                    st.warning(
                        f"⚠️ {len(code_mixed_sections)} section(s) contain code-mixed text: "
                        f"{', '.join(code_mixed_sections)}"
                    )
        
        st.markdown("---")
        
        # Multilingual summarization
        st.subheader("🤖 Multilingual Summarization")
        
        col1, col2, col3 = st.columns(3)
        
        with col1:
            output_language = st.selectbox(
                "Output Language",
                options=[
                    ('en', 'English'),
                    ('hi', 'Hindi (हिन्दी)'),
                    ('ta', 'Tamil (தமிழ்)'),
                    ('te', 'Telugu (తెలుగు)'),
                    ('bn', 'Bengali (বাংলা)'),
                    ('mr', 'Marathi (मराठी)'),
                    ('gu', 'Gujarati (ગુજરાતી)'),
                    ('kn', 'Kannada (ಕನ್ನಡ)'),
                    ('ml', 'Malayalam (മലയാളം)'),
                    ('pa', 'Punjabi (ਪੰਜਾਬੀ)')
                ],
                format_func=lambda x: x[1],
                help="Choose the language for summaries"
            )
            output_lang_code = output_language[0]
        
        with col2:
            summ_approach = st.radio(
                "Approach",
                ["Auto (Best)", "Direct", "Translate"],
                help="Auto: Choose best method automatically\\nDirect: Use multilingual model\\nTranslate: Translate → Summarize → Translate back"
            )
            
            method_map = {
                "Auto (Best)": "auto",
                "Direct": "direct",
                "Translate": "translate"
            }
            summ_method = method_map[summ_approach]
        
        with col3:
            use_gpu_multi = st.checkbox(
                "Use GPU",
                value=True,
                key="gpu_multi",
                help="Use GPU for faster processing"
            )
            
        multi_compression = st.slider(
            "Compression Ratio (% of original to keep)",
            min_value=10,
            max_value=80,
            value=30,
            step=5,
            key="multi_compression",
            help="Choose how concise the summaries should be"
        )
        
        if st.button("🚀 Generate Multilingual Summaries", type="primary"):
            try:
                with st.spinner(f"🌍 Generating summaries in {output_language[1]}..."):
                    from src.multilingual_summarizer import MultilingualSummarizer
                    
                    # Initialize
                    multi_summarizer = MultilingualSummarizer(use_gpu=use_gpu_multi)
                    
                    # Generate summaries
                    start_time = time.time()
                    results = multi_summarizer.summarize_sections(
                        sections,
                        output_lang=output_lang_code,
                        auto_detect=True,
                        method=summ_method,
                        compression=multi_compression / 100.0
                    )
                    total_time = time.time() - start_time
                    
                    # Store in session state
                    st.session_state['multi_summaries'] = results
                    st.session_state['multi_lang'] = output_lang_code
                    st.session_state['multi_time'] = total_time
                    st.success(f"✅ Summaries generated in {total_time:.1f}s!")
            except Exception as e:
                st.error(f"❌ Multilingual summarization failed: {str(e)}")
                st.info("💡 Tip: Try the 'Translate' approach if 'Direct' fails.")
        
        # Display multilingual summaries
        if 'multi_summaries' in st.session_state:
            # Check if selection matches results
            if st.session_state.get('multi_lang') != output_lang_code:
                st.warning(f"⚠️ Selected language ({output_language[1]}) differs from generated summaries ({st.session_state.get('multi_lang')}).")
                st.info("Click 'Generate Multilingual Summaries' above to update results.")
                show_results = False
            else:
                show_results = True
            
            if show_results:
                results = st.session_state['multi_summaries']
                output_lang = st.session_state['multi_lang']
                
                st.markdown("---")
            
            # Statistics
            st.subheader("📊 Summary Statistics")
            
            total_orig = sum(r.get('original_length', 0) for r in results.values())
            total_summ = sum(r.get('summary_length', 0) for r in results.values())
            avg_compression = (1 - total_summ/total_orig) * 100 if total_orig > 0 else 0
            
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.metric("Original Words", f"{total_orig:,}")
            with col2:
                st.metric("Summary Words", f"{total_summ:,}")
            with col3:
                st.metric("Compression", f"{avg_compression:.1f}%")
            with col4:
                st.metric("Time", f"{st.session_state['multi_time']:.1f}s")
            
            st.markdown("---")
            
            # Display summaries
            st.subheader(f"📝 Summaries in {output_language[1]}")
            
            for i, (section_name, result) in enumerate(results.items(), 1):
                with st.expander(
                    f"**{i}. {section_name}** "
                    f"({result.get('original_length', 0)} → {result.get('summary_length', 0)} words)",
                    expanded=(i <= 2)
                ):
                    # Metadata
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        source_lang = result.get('source_lang', 'unknown')
                        source_name = detector.get_language_info(source_lang).get('name', source_lang)
                        st.metric("Source", source_name)
                    with col2:
                        target_lang = result.get('output_lang', output_lang)
                        target_name = detector.get_language_info(target_lang).get('name', target_lang)
                        st.metric("Output", target_name)
                    with col3:
                        st.metric("Method", result.get('method', 'unknown'))
                    with col4:
                        st.metric("Time", f"{result.get('processing_time', 0):.2f}s")
                    
                    st.markdown("---")
                    
                    # Summary
                    st.markdown(f"### 📝 Summary in {output_name}")
                    
                    # Special styling based on language
                    font_size = "20px" if output_lang != 'en' else "18px"
                    
                    st.markdown(f"""
                    <div class="summary-card" style="font-size: {font_size};">
                        {result['summary']}
                    </div>
                    """, unsafe_allow_html=True)
                    
                    # Copy button in expander
                    with st.expander("📋 Copy Multilingual Summary"):
                        st.code(result['summary'], language=None)
                    
                    # Show original if different language
                    if source_lang != output_lang:
                        with st.expander("📄 View Original"):
                            original_text = sections[section_name]
                            st.text_area(
                                "Original Text",
                                original_text,
                                height=200,
                                label_visibility="collapsed",
                                key=f"multi_orig_{i}"
                            )
            
            # Download options
            st.markdown("---")
            st.subheader("📥 Download Multilingual Summaries")
            
            # Create download text
            download_text = f"""
MULTILINGUAL ANNUAL REPORT SUMMARY
{'=' * 80}

File: {uploaded_file.name}
Output Language: {output_language[1]}
Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}

Total Compression: {avg_compression:.1f}%
Processing Time: {st.session_state['multi_time']:.1f}s

{'=' * 80}

"""
            
            for section_name, result in results.items():
                download_text += f"\n{'=' * 80}\n"
                download_text += f"{section_name.upper()}\n"
                download_text += f"{'=' * 80}\n\n"
                download_text += f"Source: {result.get('source_lang', 'unknown')} | "
                download_text += f"Output: {result.get('output_lang', output_lang)} | "
                download_text += f"Method: {result.get('method', 'unknown')}\n"
                download_text += f"Compression: {result.get('compression_ratio', 0)*100:.1f}%\n\n"
                download_text += result['summary'] + "\n"
            
            st.download_button(
                label=f"📥 Download All ({output_language[1]})",
                data=download_text.encode('utf-8'),
                file_name=f"{Path(uploaded_file.name).stem}_summaries_{output_lang}.txt",
                mime="text/plain"
            )
        
        else:
            st.info("👆 Select output language and click 'Generate' to create multilingual summaries")
            
            # Show supported languages
            st.markdown("### 🗣️ Supported Languages")
            
            supported = [
                "🇮🇳 Hindi (हिन्दी)",
                "🇮🇳 Tamil (தமிழ்)",
                "🇮🇳 Telugu (తెలుగు)",
                "🇮🇳 Bengali (বাংলা)",
                "🇮🇳 Marathi (मराठी)",
                "🇮🇳 Gujarati (ગુજરાતી)",
                "🇮🇳 Kannada (ಕನ್ನಡ)",
                "🇮🇳 Malayalam (മലയാളം)",
                "🇮🇳 Punjabi (ਪੰਜਾਬੀ)",
                "🇮🇳 Odia (ଓଡ଼ିଆ)",
                "🇮🇳 Assamese (অসমীয়া)",
                "🇬🇧 English"
            ]
            
            cols = st.columns(3)
            for i, lang in enumerate(supported):
                with cols[i % 3]:
                    st.markdown(f"✓ {lang}")
        
        # Additional info
        st.markdown("---")
        st.info("""
        **💡 How It Works:**
        
        - **Auto Mode**: Detects source language and uses the best method
        - **Direct Mode**: Uses IndicBART (AI4Bharat model) for direct Indian language summarization
        - **Translate Mode**: Translates to English → Summarizes → Translates back
        
        **Note**: First use will download multilingual models (~1-2GB)
        """)

    with tabs[7]:  # Download tab (was 6)
        st.header("Download Results")

        # Save sections
        output_dir = processor.save_sections()

        st.success("✅ Sections saved to files!")
        st.info(f"📁 Location: `{output_dir}`")

        # Create downloadable summary
        summary_text = f"""ANNUAL REPORT ANALYSIS
{'=' * 50}

File: {uploaded_file.name}
Pages: {processor.metadata.get('total_pages', 0)}
Sections: {len(sections)}

{'=' * 50}
SECTIONS SUMMARY
{'=' * 50}

"""

        for section_name, content in sections.items():
            summary_text += f"\n\n{'=' * 50}\n"
            summary_text += f"{section_name.upper()}\n"
            summary_text += f"{'=' * 50}\n\n"
            summary_text += content[:500] + "...\n"

        # Download button
        st.download_button(
            label="📥 Download Complete Analysis",
            data=summary_text,
            file_name=f"{Path(uploaded_file.name).stem}_analysis.txt",
            mime="text/plain"
        )
        
        # DOCX Export
        def create_docx(sections):
            try:
                from docx import Document
                doc = Document()
                doc.add_heading('Annual Report Analysis', 0)
                
                for section, content in sections.items():
                    doc.add_heading(section, level=1)
                    doc.add_paragraph(content[:2000] + "..." if len(content) > 2000 else content)
                
                bio = io.BytesIO()
                doc.save(bio)
                return bio.getvalue()
            except Exception as e:
                st.error(f"DOCX generation failed: {e}")
                return None
            
        docx_data = create_docx(sections)
        if docx_data:
            st.download_button(
                label="📥 Download as Word (DOCX)",
                data=docx_data,
                file_name=f"{Path(uploaded_file.name).stem}_analysis.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

        # Individual section downloads
        st.markdown("---")
        st.subheader("Download Individual Sections")

        cols = st.columns(3)
        for i, (section_name, content) in enumerate(sections.items()):
            with cols[i % 3]:
                st.download_button(
                    label=f"📄 {section_name}",
                    data=content,
                    file_name=f"{section_name.replace(' ', '_')}.txt",
                    mime="text/plain",
                    key=f"download_{i}"
                )

    # Results persist in session state
    # We remove the unlink to ensure the PDF remains available for visualizations
    pass

else:
    # Landing page when no file is uploaded
    st.info("👆 Upload a PDF file to get started!")

    st.markdown("---")
    st.subheader("✨ What This Tool Does")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown("""
        ### 📄 Text Extraction
        Extracts all text from PDF while preserving page structure and formatting.
        """)

    with col2:
        st.markdown("""
        ### 🔍 Section Detection
        Automatically identifies key sections like Financial Performance, Chairman's Message, etc.
        """)

    with col3:
        st.markdown("""
        ### 📊 Data Extraction
        Extracts tables and statistical information from the document.
        """)

    st.markdown("---")
    st.subheader("📋 Supported Sections")

    sections_info = {
        "Chairman's Message": "Leadership communication to shareholders",
        "Executive Summary": "Key highlights and overview",
        "Financial Performance": "Revenue, profit, and financial metrics",
        "Management Discussion": "Detailed analysis by management",
        "Business Overview": "Company operations and segments",
        "Corporate Governance": "Governance practices and policies",
        "Risk Management": "Risk factors and mitigation strategies",
        "Future Outlook": "Forward-looking statements and plans",
        "Financial Statements": "Balance sheet, income statement, etc.",
        "Auditor's Report": "Independent auditor's opinion"
    }

    for section, description in sections_info.items():
        st.markdown(f"**{section}**: {description}")

# Footer
st.markdown("---")
st.markdown("""
<div style='text-align: center; color: #666;'>
    <p>📊 Annual Report Analyzer | Day 5/6 Progress | NLP Project</p>
</div>
""", unsafe_allow_html=True)
